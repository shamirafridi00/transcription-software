import os
from flask import Flask, render_template, request, send_file
import whisper
from docx import Document
import subprocess
import torch

app = Flask(__name__)

# Configure Whisper for optimal CPU performance
model = None

def load_whisper_model():
    global model
    if model is None:
        # Load the smallest English-only model to save memory
        model = whisper.load_model("tiny.en", device="cpu")
    return model

def convert_ts_to_wav(ts_file, wav_file):
    """Convert TS file to WAV using ffmpeg"""
    command = [
        'ffmpeg', '-i', ts_file,
        '-ac', '1',  # Convert to mono
        '-ar', '16000',  # Sample rate 16kHz
        '-y',  # Overwrite output file if it exists
        wav_file
    ]
    subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

def process_audio(file_path):
    """Process single audio file and return transcription"""
    try:
        model = load_whisper_model()
        
        # Convert TS to WAV if needed
        if file_path.endswith('.ts'):
            wav_path = file_path.replace('.ts', '.wav')
            convert_ts_to_wav(file_path, wav_path)
            result = model.transcribe(wav_path, language='en')
            if os.path.exists(wav_path):
                os.remove(wav_path)  # Clean up temporary WAV file
        else:
            result = model.transcribe(file_path, language='en')
        
        return result["text"]
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")
        return f"Error processing file: {str(e)}"

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        files = request.files.getlist('files')
        
        if not files[0].filename:
            return 'No files selected'

        # Create a folder for outputs if it doesn't exist
        if not os.path.exists('outputs'):
            os.makedirs('outputs')

        processed_files = []
        # Process each file
        for file in files:
            if file.filename.endswith('.ts'):
                try:
                    # Save uploaded file temporarily
                    temp_path = os.path.join('outputs', file.filename)
                    file.save(temp_path)

                    # Create document
                    doc = Document()
                    doc.add_heading(f'Transcription - {file.filename}', 0)

                    # Process and add transcription
                    transcription = process_audio(temp_path)
                    doc.add_paragraph(transcription)

                    # Save document
                    docx_path = os.path.join('outputs', f'{file.filename[:-3]}.docx')
                    doc.save(docx_path)

                    # Clean up temporary TS file
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                    
                    processed_files.append(file.filename)
                except Exception as e:
                    return f'Error processing {file.filename}: {str(e)}'

        return f'Successfully processed files: {", ".join(processed_files)}. Check the outputs folder.'

    return '''
    <!doctype html>
    <html>
    <head>
        <title>TS to DOCX Converter</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 40px; background-color: #f5f5f5; }
            .container { max-width: 600px; margin: 0 auto; }
            .upload-form { 
                padding: 20px; 
                border: 1px solid #ccc; 
                border-radius: 5px;
                background-color: white;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }
            .submit-btn { 
                margin-top: 10px; 
                padding: 10px 20px;
                background-color: #4CAF50;
                color: white;
                border: none;
                border-radius: 4px;
                cursor: pointer;
            }
            .submit-btn:hover {
                background-color: #45a049;
            }
            input[type="file"] {
                margin: 10px 0;
                padding: 10px;
                width: 100%;
                box-sizing: border-box;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>TS to DOCX Converter</h1>
            <div class="upload-form">
                <form method="post" enctype="multipart/form-data">
                    <p>Select .ts files to convert:</p>
                    <input type="file" name="files" multiple accept=".ts">
                    <br>
                    <input type="submit" value="Convert Files" class="submit-btn">
                </form>
            </div>
        </div>
    </body>
    </html>
    '''

if __name__ == '__main__':
    app.run(debug=True, port=5000)